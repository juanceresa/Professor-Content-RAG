import os
import logging
from pathlib import Path
from typing import List, Dict, Optional
import hashlib
from dataclasses import dataclass
from datetime import datetime

# Document processing
from docx import Document
import PyPDF2
import pdfplumber
from sentence_transformers import SentenceTransformer

# Vector database
from pinecone import Pinecone, ServerlessSpec
import numpy as np

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

@dataclass
class DocumentChunk:
    """Represents a chunk of processed document content"""
    text: str
    metadata: Dict
    chunk_id: str
    embedding: Optional[List[float]] = None

class DocumentProcessor:
    """Handles processing of course documents into embeddings"""

    def __init__(self, embedding_model_name: str = "sentence-transformers/all-mpnet-base-v2"):
        self.embedding_model = SentenceTransformer(embedding_model_name)
        logger.info(f"Initialized embedding model: {embedding_model_name}")

    def extract_text_from_docx(self, file_path: Path) -> str:
        """Extract text from Word document"""
        try:
            doc = Document(file_path)
            full_text = []

            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    full_text.append(paragraph.text.strip())

            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        full_text.append(" | ".join(row_text))

            return "\n\n".join(full_text)

        except Exception as e:
            logger.error(f"Error extracting text from {file_path}: {e}")
            return ""

    def extract_text_from_pdf(self, file_path: Path) -> str:
        """Extract text from PDF document preserving hierarchical structure"""
        try:
            text_blocks = []
            
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    # Extract text with layout preservation - key for hierarchical bullets
                    page_text = page.extract_text(layout=True, x_tolerance=3, y_tolerance=3)
                    
                    if page_text:
                        # Split into lines and preserve indentation
                        lines = page_text.split('\n')
                        for line in lines:
                            if line.strip():  # Skip empty lines
                                text_blocks.append(line.rstrip())
            
            extracted_text = "\n".join(text_blocks)
            logger.info(f"pdfplumber extracted {len(extracted_text)} characters from {file_path.name}")
            return extracted_text

        except Exception as e:
            logger.error(f"Error extracting text with pdfplumber from {file_path}: {e}")
            
            # Fallback to PyPDF2 if pdfplumber fails
            try:
                import PyPDF2
                text = []
                with open(file_path, 'rb') as file:
                    pdf_reader = PyPDF2.PdfReader(file)
                    for page in pdf_reader.pages:
                        page_text = page.extract_text()
                        if page_text.strip():
                            text.append(page_text.strip())
                logger.info(f"PyPDF2 fallback extraction from {file_path.name}")
                return "\n\n".join(text)
            except Exception as e2:
                logger.error(f"Error with PyPDF2 fallback for {file_path}: {e2}")
                return ""

        except Exception as e:
            logger.error(f"Error extracting text from {file_path}: {e}")
            return ""

    def extract_text_from_txt(self, file_path: Path) -> str:
        """Extract text from plain text file"""
        try:
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as file:
                return file.read()
        except Exception as e:
            logger.error(f"Error reading text file {file_path}: {e}")
            return ""

    def extract_text_from_file(self, file_path: Path) -> str:
        """Extract text from various file formats"""
        suffix = file_path.suffix.lower()

        if suffix in ['.docx', '.doc']:
            return self.extract_text_from_docx(file_path)
        elif suffix == '.pdf':
            return self.extract_text_from_pdf(file_path)
        elif suffix in ['.txt', '.md']:
            return self.extract_text_from_txt(file_path)
        else:
            logger.warning(f"Unsupported file format: {suffix}")
            return ""

    def extract_lesson_structure(self, text: str, doc_type: str) -> Dict:
        """Extract lessons with 100% confidence in MASTER files"""
        import re
        lessons = {}
        
        if doc_type == "master_lecture":
            lines = text.split('\n')
            for i, line in enumerate(lines):
                stripped = line.strip()
                
                # Lesson headers: "Lesson X" with no bullet point
                lesson_match = re.match(r'^Lesson\s+(\d+)(?:\s|$)', stripped, re.IGNORECASE)
                if lesson_match and not any(stripped.startswith(bullet) for bullet in ['•', '○', '■', '▪', '-', '*']):
                    lesson_number = lesson_match.group(1)
                    lessons[lesson_number] = {
                        'lesson_number': lesson_number,
                        'content_start_line': i,
                        'title': stripped,
                        'raw_line': line
                    }
                    logger.info(f"Found lesson {lesson_number}: {stripped}")
        
        return lessons

    def extract_hierarchical_structure(self, text: str, doc_type: str = "general") -> List[Dict]:
        """Extract hierarchical bullet structure from academic content with enhanced lesson detection"""
        import re
        lines = text.split('\n')
        structured_content = []
        current_lesson = None
        hierarchy_stack = []

        # Bullet symbols in order of hierarchy depth
        bullet_patterns = [
            ('•', 1),  # Main topic bullet
            ('○', 2),  # Subtopic bullet
            ('■', 3),  # Detail bullet
            ('▪', 4),  # Sub-detail bullet
            ('-', 2),  # Alternative subtopic
            ('*', 2),  # Alternative subtopic
        ]

        for line in lines:
            stripped_line = line.strip()
            if not stripped_line:
                continue

            # Enhanced lesson detection for MASTER lecture files
            is_lesson_header = False
            if doc_type == "master_lecture":
                # Look for "Lesson X" pattern with no bullet points
                lesson_match = re.match(r'^Lesson\s+(\d+)(?:\s|$)', stripped_line, re.IGNORECASE)
                if lesson_match and not any(stripped_line.startswith(bullet) for bullet in ['•', '○', '■', '▪', '-', '*']):
                    current_lesson = stripped_line
                    hierarchy_stack = [current_lesson]
                    is_lesson_header = True
                    
                    # Add lesson header as a structured item
                    structured_content.append({
                        'content': stripped_line,
                        'level': 0,  # Lesson level
                        'symbol': None,
                        'hierarchy_path': current_lesson,
                        'lesson': current_lesson,
                        'full_line': line,
                        'is_lesson_header': True
                    })
                    continue
            else:
                # General lesson detection for other documents
                if 'lesson' in stripped_line.lower() and any(word in stripped_line.lower() for word in ['lesson', 'chapter', 'module']):
                    current_lesson = stripped_line
                    hierarchy_stack = [current_lesson]
                    is_lesson_header = True
                    continue

            # Detect bullet points and their hierarchy level
            bullet_level = 0
            bullet_symbol = None
            content = stripped_line

            for symbol, level in bullet_patterns:
                if stripped_line.startswith(symbol + ' '):
                    bullet_level = level
                    bullet_symbol = symbol
                    content = stripped_line[len(symbol):].strip()
                    break

            # Calculate indentation level as fallback
            if bullet_level == 0:
                indent_level = len(line) - len(line.lstrip())
                if indent_level > 0:
                    bullet_level = min(4, (indent_level // 2) + 1)

            # Update hierarchy stack
            if bullet_level > 0:
                # Trim stack to current level
                hierarchy_stack = hierarchy_stack[:bullet_level]
                hierarchy_stack.append(content)

                # Create hierarchy path
                hierarchy_path = " > ".join(hierarchy_stack)

                structured_content.append({
                    'content': content,
                    'level': bullet_level,
                    'symbol': bullet_symbol,
                    'hierarchy_path': hierarchy_path,
                    'lesson': current_lesson,
                    'full_line': line
                })
            else:
                # Regular paragraph content
                if hierarchy_stack:
                    hierarchy_path = " > ".join(hierarchy_stack)
                else:
                    hierarchy_path = current_lesson or "General Content"

                structured_content.append({
                    'content': content,
                    'level': 0,
                    'symbol': None,
                    'hierarchy_path': hierarchy_path,
                    'lesson': current_lesson,
                    'full_line': line
                })

        return structured_content

    def hierarchical_chunking(self, structured_content: List[Dict], doc_type: str = "general", target_size: int = 900, overlap_size: int = 150) -> List[str]:
        """Create chunks that preserve hierarchical structure and conceptual flow"""
        if not structured_content:
            return []

        if doc_type == "master_lecture":
            return self._create_lesson_based_chunks(structured_content, target_size)
        else:
            return self._create_topic_based_chunks(structured_content, target_size)

    def _create_lesson_based_chunks(self, content: List[Dict], target_size: int = 900) -> List[str]:
        """Chunk MASTER PDFs with lesson boundaries as primary breaks"""
        chunks = []
        current_lesson_content = []
        current_lesson = None
        
        for item in content:
            # Hard boundary at lesson changes
            if item.get('lesson') != current_lesson and current_lesson_content:
                # Process the completed lesson
                lesson_chunks = self._chunk_within_lesson(current_lesson_content, target_size)
                chunks.extend(lesson_chunks)
                current_lesson_content = []
            
            current_lesson = item.get('lesson')
            current_lesson_content.append(item)
        
        # Process final lesson
        if current_lesson_content:
            lesson_chunks = self._chunk_within_lesson(current_lesson_content, target_size)
            chunks.extend(lesson_chunks)
        
        return chunks

    def _chunk_within_lesson(self, lesson_content: List[Dict], target_size: int = 900) -> List[str]:
        """Chunk within a lesson, respecting bullet hierarchy"""
        chunks = []
        current_chunk_items = []
        current_size = 0
        
        i = 0
        while i < len(lesson_content):
            item = lesson_content[i]
            
            # Get the complete bullet tree starting from this item
            bullet_tree = self._extract_complete_bullet_tree(lesson_content, i)
            tree_size = sum(len(item['full_line']) for item in bullet_tree)
            
            # If adding this tree exceeds size AND we have content, create chunk
            if current_size + tree_size > target_size and current_chunk_items:
                chunk = self._assemble_chunk_with_context(current_chunk_items)
                chunks.append(chunk)
                current_chunk_items = []
                current_size = 0
            
            # Add the complete bullet tree
            current_chunk_items.extend(bullet_tree)
            current_size += tree_size
            
            # Skip ahead past the tree items we just processed
            i += len(bullet_tree)
        
        # Final chunk
        if current_chunk_items:
            chunk = self._assemble_chunk_with_context(current_chunk_items)
            chunks.append(chunk)
        
        return chunks

    def _extract_complete_bullet_tree(self, content: List[Dict], start_idx: int) -> List[Dict]:
        """Extract a complete bullet tree to avoid breaking conceptual flow"""
        if start_idx >= len(content):
            return []
        
        tree = [content[start_idx]]
        start_item = content[start_idx]
        start_level = start_item.get('level', 0)
        
        # If this is a lesson header, just return it
        if start_item.get('is_lesson_header'):
            return tree
        
        # If this is a main bullet (•), include all its sub-bullets
        if start_level == 1:  # Main bullet
            i = start_idx + 1
            while i < len(content):
                item = content[i]
                item_level = item.get('level', 0)
                
                # Stop when we hit another main bullet, lesson header, or lesson boundary
                if ((item_level <= 1 and i > start_idx) or 
                    item.get('is_lesson_header') or
                    item.get('lesson') != start_item.get('lesson')):
                    break
                
                tree.append(item)
                i += 1
        
        return tree

    def _create_topic_based_chunks(self, content: List[Dict], target_size: int = 900) -> List[str]:
        """Create topic-based chunks for reading materials"""
        # Similar to lesson-based but without lesson boundaries
        chunks = []
        current_chunk_items = []
        current_size = 0
        
        i = 0
        while i < len(content):
            item = content[i]
            
            # Get bullet tree for main bullets, single item for others
            if item.get('level') == 1:  # Main bullet
                bullet_tree = self._extract_complete_bullet_tree(content, i)
            else:
                bullet_tree = [item]
            
            tree_size = sum(len(item['full_line']) for item in bullet_tree)
            
            # Check size constraints
            if current_size + tree_size > target_size and current_chunk_items:
                chunk = self._assemble_chunk_with_context(current_chunk_items)
                chunks.append(chunk)
                current_chunk_items = []
                current_size = 0
            
            current_chunk_items.extend(bullet_tree)
            current_size += tree_size
            i += len(bullet_tree)
        
        # Final chunk
        if current_chunk_items:
            chunk = self._assemble_chunk_with_context(current_chunk_items)
            chunks.append(chunk)
        
        return chunks

    def _assemble_chunk_with_context(self, items: List[Dict]) -> str:
        """Assemble chunk with proper hierarchical context and lesson information"""
        if not items:
            return ""
        
        chunk_lines = []
        
        # Add lesson header if present
        lesson = items[0].get('lesson')
        if lesson and not items[0].get('is_lesson_header'):
            chunk_lines.append(f"=== {lesson} ===\n")
        
        # Add items preserving original formatting and indentation
        for item in items:
            chunk_lines.append(item['full_line'])
        
        return '\n'.join(chunk_lines)

    def _create_chunk_with_context(self, items: List[Dict]) -> str:
        """Legacy method - redirects to new assembly method"""
        return self._assemble_chunk_with_context(items)

    def _get_overlap_items(self, items: List[Dict], overlap_size: int) -> List[Dict]:
        """Get items for overlap, prioritizing higher-level hierarchy items"""
        if not items:
            return []

        # Start from the end and work backwards, prioritizing higher hierarchy levels
        overlap_items = []
        current_size = 0

        for item in reversed(items):
            item_size = len(item['full_line'])
            if current_size + item_size <= overlap_size:
                overlap_items.insert(0, item)
                current_size += item_size
            elif item['level'] <= 2:  # Always include high-level context
                overlap_items.insert(0, item)
                break

        return overlap_items

    def intelligent_chunking(self, text: str, doc_type: str = "general", max_chunk_size: int = 900, overlap: int = 150) -> List[str]:
        """Enhanced chunking that preserves hierarchical academic structure"""
        if not text.strip():
            return []

        # First, try hierarchical chunking for structured content
        structured_content = self.extract_hierarchical_structure(text, doc_type)

        if len(structured_content) > 3:  # If we found good hierarchical structure
            return self.hierarchical_chunking(structured_content, doc_type, target_size=max_chunk_size, overlap_size=overlap)

        # Fallback to paragraph-based chunking for unstructured content
        paragraphs = [p.strip() for p in text.split('\n\n') if p.strip()]

        chunks = []
        current_chunk = ""

        for paragraph in paragraphs:
            # If adding this paragraph would exceed max size, save current chunk
            if len(current_chunk) + len(paragraph) > max_chunk_size and current_chunk:
                chunks.append(current_chunk.strip())

                # Start new chunk with overlap from previous chunk
                if overlap > 0:
                    words = current_chunk.split()
                    overlap_text = " ".join(words[-overlap//4:])  # Last ~25% of words for overlap
                    current_chunk = overlap_text + "\n\n" + paragraph
                else:
                    current_chunk = paragraph
            else:
                if current_chunk:
                    current_chunk += "\n\n" + paragraph
                else:
                    current_chunk = paragraph

        # Add the last chunk
        if current_chunk.strip():
            chunks.append(current_chunk.strip())

        return chunks

    def infer_document_type(self, file_path: Path) -> str:
        """Enhanced document classification with lesson detection priority"""
        filename = file_path.name
        
        # MASTER lecture PDFs - highest priority for lesson extraction
        if (filename.startswith("MASTER") and 
            "Lecture" in filename and 
            filename.endswith(".pdf")):
            return "master_lecture"
        
        # MASTER reading PDFs - reading materials in PDF format
        elif (filename.startswith("MASTER") and 
              "Read" in filename and 
              filename.endswith(".pdf")):
            return "reading_material"
        
        # Reading materials - support content without lesson structure
        elif file_path.suffix.lower() in ['.docx', '.doc']:
            return "reading_material"
        
        # Legacy classification for other files
        filename_lower = filename.lower()
        if 'syllabus' in filename_lower:
            return 'syllabus'
        elif 'lecture' in filename_lower or 'slides' in filename_lower:
            return 'lecture'
        elif 'reading' in filename_lower or 'assignment' in filename_lower:
            return 'reading_material'
        elif 'exam' in filename_lower or 'quiz' in filename_lower:
            return 'assessment'
        elif filename.endswith(".pdf"):
            return "supplementary_pdf"
        else:
            return 'general'

    def process_document(self, file_path: Path, course_name: str, professor_name: str) -> List[DocumentChunk]:
        """Process a single document into chunks with metadata"""
        logger.info(f"Processing document: {file_path}")

        # Extract text
        text = self.extract_text_from_file(file_path)
        if not text.strip():
            logger.warning(f"No text extracted from {file_path}")
            return []

        # Get document type for enhanced processing
        doc_type = self.infer_document_type(file_path)
        
        # Import improved chunking algorithm
        import sys
        from pathlib import Path
        
        # Add parent directory to path for improved_chunking import
        parent_dir = Path(__file__).parent.parent
        if str(parent_dir) not in sys.path:
            sys.path.insert(0, str(parent_dir))
        
        from improved_chunking import create_improved_chunks
        
        # Extract lesson number if available
        lesson_number = None
        if 'lesson' in file_path.name.lower():
            import re
            match = re.search(r'lesson(\d+(?:\.\d+)?)', file_path.name.lower())
            if match:
                lesson_number = match.group(1)
        
        # Create chunks using improved algorithm
        chunks = create_improved_chunks(
            text=text, 
            lesson_number=lesson_number,
            min_size=300,     # Smaller minimum for academic content
            target_size=700,  # Good balance for retrieval
            max_size=1000     # Not too large for context windows
        )
        logger.info(f"Created {len(chunks)} improved chunks from {file_path}")

        # Create DocumentChunk objects with enhanced metadata
        document_chunks = []
        for i, chunk_text in enumerate(chunks):
            # Create unique chunk ID
            chunk_content = f"{file_path}_{i}_{chunk_text[:50]}"
            chunk_id = hashlib.md5(chunk_content.encode()).hexdigest()

            # Extract hierarchy information from chunk
            hierarchy_info = self._extract_chunk_hierarchy(chunk_text)

            # Create namespace for consistency with Pinecone storage
            # Find the course key that matches this course name
            course_mapping = get_course_mapping()
            namespace = None
            for key, config in course_mapping.items():
                if config['name'] == course_name:
                    namespace = key
                    break

            if not namespace:
                # Fallback: create namespace from course name
                namespace = course_name.lower().replace(" ", "_").replace("-", "_")

            # Enhanced hierarchy extraction with document type awareness
            hierarchy_info = self._extract_chunk_hierarchy(chunk_text, doc_type, file_path)

            metadata = {
                'course': namespace,  # Use namespace for consistency with search filter
                'course_name': course_name,  # Keep original course name for display
                'professor': professor_name,
                'document_name': file_path.name,
                'document_type': doc_type,
                'chunk_index': i,
                'total_chunks': len(chunks),
                'file_path': str(file_path),
                'processed_date': datetime.now().isoformat(),
                'text': chunk_text,  # Store text in metadata for Pinecone

                # Enhanced hierarchical metadata (ensure no null values)
                'lesson_number': hierarchy_info.get('lesson_number') or "",
                'main_topic': hierarchy_info.get('main_topic') or "",
                'hierarchy_path': hierarchy_info.get('hierarchy_path') or "General Content",
                'depth_level': hierarchy_info.get('depth_level', 0),
                'chunk_type': hierarchy_info.get('chunk_type', 'general'),
                'conceptual_completeness': hierarchy_info.get('conceptual_completeness', 'complete')
            }

            # Generate embedding for this chunk
            embedding = self.embedding_model.encode([chunk_text])[0].tolist()

            document_chunks.append(DocumentChunk(
                text=chunk_text,
                metadata=metadata,
                chunk_id=chunk_id,
                embedding=embedding
            ))

        return document_chunks

    def _extract_chunk_hierarchy(self, chunk_text: str, doc_type: str = "general", file_path: Path = None) -> Dict:
        """Extract enhanced hierarchical information from a chunk with document type awareness"""
        import re
        lines = chunk_text.split('\n')
        hierarchy_info = {
            'lesson_number': None,
            'main_topic': None,
            'hierarchy_path': None,
            'depth_level': 0,
            'chunk_type': 'general',
            'conceptual_completeness': 'complete'
        }

        # Enhanced lesson detection based on document type
        for line in lines:
            stripped_line = line.strip()

            # Enhanced lesson number extraction
            if not hierarchy_info['lesson_number']:
                # Look for lesson markers (=== Lesson X ===)
                lesson_marker_match = re.search(r'=== (Lesson \d+) ===', stripped_line)
                if lesson_marker_match:
                    lesson_num_match = re.search(r'Lesson (\d+)', lesson_marker_match.group(1))
                    if lesson_num_match:
                        hierarchy_info['lesson_number'] = lesson_num_match.group(1)
                        continue

                # Look for direct lesson headers
                elif doc_type == "master_lecture":
                    lesson_match = re.match(r'^Lesson\s+(\d+)(?:\s|$)', stripped_line, re.IGNORECASE)
                    if lesson_match and not any(stripped_line.startswith(bullet) for bullet in ['•', '○', '■', '▪', '-', '*']):
                        hierarchy_info['lesson_number'] = lesson_match.group(1)
                        continue

                # General lesson detection for other documents
                elif 'lesson' in stripped_line.lower():
                    lesson_match = re.search(r'lesson\s+(\d+)', stripped_line.lower())
                    if lesson_match:
                        hierarchy_info['lesson_number'] = lesson_match.group(1)

            # Detect main topic (usually the first major bullet point)
            if stripped_line.startswith('•') and not hierarchy_info['main_topic']:
                hierarchy_info['main_topic'] = stripped_line[1:].strip()

            # Calculate maximum depth level in chunk
            if any(stripped_line.startswith(symbol) for symbol in ['•', '○', '■', '▪']):
                if stripped_line.startswith('•'):
                    hierarchy_info['depth_level'] = max(hierarchy_info['depth_level'], 1)
                elif stripped_line.startswith('○'):
                    hierarchy_info['depth_level'] = max(hierarchy_info['depth_level'], 2)
                elif stripped_line.startswith('■'):
                    hierarchy_info['depth_level'] = max(hierarchy_info['depth_level'], 3)
                elif stripped_line.startswith('▪'):
                    hierarchy_info['depth_level'] = max(hierarchy_info['depth_level'], 4)

        # Determine chunk type
        if hierarchy_info['lesson_number']:
            hierarchy_info['chunk_type'] = 'lesson_content'
        elif hierarchy_info['depth_level'] > 0:
            hierarchy_info['chunk_type'] = 'structured_content'

        # Enhanced hierarchy path building
        path_components = []
        if hierarchy_info['lesson_number']:
            path_components.append(f"Lesson {hierarchy_info['lesson_number']}")
        if hierarchy_info['main_topic']:
            path_components.append(hierarchy_info['main_topic'][:50])  # Truncate if too long

        hierarchy_info['hierarchy_path'] = " > ".join(path_components) if path_components else "General Content"

        # Determine conceptual completeness based on chunk structure
        if hierarchy_info['main_topic'] and hierarchy_info['depth_level'] > 1:
            hierarchy_info['conceptual_completeness'] = 'complete'
        elif hierarchy_info['main_topic']:
            hierarchy_info['conceptual_completeness'] = 'partial'
        else:
            hierarchy_info['conceptual_completeness'] = 'fragment'

        # Set chunk type based on document type and content
        if doc_type == "master_lecture" and hierarchy_info['lesson_number']:
            hierarchy_info['chunk_type'] = 'lesson_content'
        elif doc_type == "reading_material":
            hierarchy_info['chunk_type'] = 'reading_content'
        elif hierarchy_info['depth_level'] > 0:
            hierarchy_info['chunk_type'] = 'structured_content'

        return hierarchy_info

    def generate_embeddings(self, chunks: List[DocumentChunk]) -> List[DocumentChunk]:
        """Generate embeddings for document chunks"""
        if not chunks:
            return chunks

        logger.info(f"Generating embeddings for {len(chunks)} chunks")

        # Extract text for embedding
        texts = [chunk.text for chunk in chunks]

        # Generate embeddings
        embeddings = self.embedding_model.encode(texts, show_progress_bar=True)

        # Add embeddings to chunks
        for chunk, embedding in zip(chunks, embeddings):
            chunk.embedding = embedding.tolist()

        return chunks

class PineconeUploader:
    """Handles uploading processed documents to Pinecone"""

    def __init__(self, api_key: str, index_name: str, dimension: int = 768):
        self.pc = Pinecone(api_key=api_key)
        self.index_name = index_name
        self.dimension = dimension

        # Create index if it doesn't exist
        self._ensure_index_exists()
        self.index = self.pc.Index(index_name)
        logger.info(f"Connected to Pinecone index: {index_name}")

    def _ensure_index_exists(self):
        """Create Pinecone index if it doesn't exist"""
        existing_indexes = [index.name for index in self.pc.list_indexes()]

        if self.index_name not in existing_indexes:
            logger.info(f"Creating new Pinecone index: {self.index_name}")
            self.pc.create_index(
                name=self.index_name,
                dimension=self.dimension,
                metric="cosine",
                spec=ServerlessSpec(cloud="aws", region="us-east-1")
            )
        else:
            logger.info(f"Using existing Pinecone index: {self.index_name}")

    def upload_chunks(self, chunks: List[DocumentChunk], namespace: str, batch_size: int = 100):
        """Upload chunks to Pinecone with specified namespace"""
        if not chunks:
            logger.warning("No chunks to upload")
            return

        logger.info(f"Uploading {len(chunks)} chunks to namespace: {namespace}")

        # Prepare vectors for upsert
        vectors = []
        for chunk in chunks:
            if chunk.embedding is None:
                logger.warning(f"Skipping chunk without embedding: {chunk.chunk_id}")
                continue

            vectors.append({
                "id": chunk.chunk_id,
                "values": chunk.embedding,
                "metadata": chunk.metadata
            })

        # Upload in batches
        for i in range(0, len(vectors), batch_size):
            batch = vectors[i:i + batch_size]
            try:
                self.index.upsert(vectors=batch, namespace=namespace)
                logger.info(f"Uploaded batch {i//batch_size + 1}/{(len(vectors)-1)//batch_size + 1}")
            except Exception as e:
                logger.error(f"Error uploading batch {i//batch_size + 1}: {e}")

        logger.info(f"Successfully uploaded {len(vectors)} chunks to namespace: {namespace}")

    def get_index_stats(self) -> Dict:
        """Get statistics about the Pinecone index"""
        return self.index.describe_index_stats()

def get_course_mapping():
    """Define how document folders map to chatbot courses"""
    return {
        "federal_state_local": {
            "name": "Federal, State, and Local Government",
            "folders": ["govt", "local"],
            "description": "American government systems at federal, state, and local levels",
            "system_prompt": """You are Professor Robert Ceresa, teaching American Government.
            You specialize in federal, state, and local government systems, institutions, and processes.
            Answer questions based on the course materials provided. Be scholarly but accessible to undergraduate students.
            If you don't have specific information from the course materials, say so clearly."""
        },
        "american": {
            "name": "American Political System",
            "folders": ["american"],
            "description": "American political institutions, processes, and governance",
            "system_prompt": """You are Professor Robert Ceresa, teaching American Politics.
            Answer questions based on the course materials provided. Be scholarly but accessible to undergraduate students.
            If you don't have specific information from the course materials, say so clearly."""
        },
        "foundational": {
            "name": "Foundational Political Theory",
            "folders": ["foundational"],
            "description": "Core concepts and foundations of political science",
            "system_prompt": """You are Professor Robert Ceresa, teaching Foundational Political Theory.
            Answer questions based on the course materials provided. Be scholarly but accessible to undergraduate students.
            If you don't have specific information from the course materials, say so clearly."""
        },
        "functional": {
            "name": "Functional Political Analysis",
            "folders": ["functional"],
            "description": "Functional approaches to understanding political systems",
            "system_prompt": """You are Professor Robert Ceresa, teaching Functional Political Analysis.
            Answer questions based on the course materials provided. Be scholarly but accessible to undergraduate students.
            If you don't have specific information from the course materials, say so clearly."""
        },
        "international": {
            "name": "International Relations & Comparative Politics",
            "folders": ["international"],
            "description": "International relations, comparative politics, and global affairs",
            "system_prompt": """You are Professor Robert Ceresa, teaching International Relations and Comparative Politics.
            Answer questions based on the course materials provided. Be scholarly but accessible to undergraduate students.
            If you don't have specific information from the course materials, say so clearly."""
        },
        "professional": {
            "name": "Professional & Management Politics",
            "folders": ["professional"],
            "description": "Professional development and management in political contexts",
            "system_prompt": """You are Professor Robert Ceresa, teaching Professional and Management Politics.
            Answer questions based on the course materials provided. Be scholarly but accessible to undergraduate students.
            If you don't have specific information from the course materials, say so clearly."""
        },
        "theory": {
            "name": "Political Philosophy & Theory",
            "folders": ["theory"],
            "description": "Classical and modern political philosophy and theory",
            "system_prompt": """You are Professor Robert Ceresa, teaching Political Philosophy and Theory.
            You specialize in classical and contemporary political thought, from Aristotle and Plato to modern theorists.
            Answer questions based on the course materials provided. Be scholarly but accessible to undergraduate students.
            If you don't have specific information from the course materials, say so clearly."""
        }
    }

def process_course_content(
    data_directory: Path,
    pinecone_api_key: str,
    pinecone_index_name: str,
    professor_name: str = "Professor"
):
    """Main function to process all courses from local data directory structure"""
    data_path = Path(data_directory)

    if not data_path.exists():
        logger.error(f"Data directory not found: {data_path}")
        return

    logger.info(f"Starting processing from data directory: {data_path}")

    # Initialize processors
    doc_processor = DocumentProcessor()
    uploader = PineconeUploader(pinecone_api_key, pinecone_index_name)

    # Get course mapping configuration
    course_mapping = get_course_mapping()

    # Find all available folders
    available_folders = [folder.name for folder in data_path.iterdir()
                        if folder.is_dir() and not folder.name.startswith('.')]

    logger.info(f"Found {len(available_folders)} document folders: {available_folders}")

    # Process each configured course
    total_chunks_processed = 0

    for course_key, course_config in course_mapping.items():
        logger.info(f"\n{'='*50}")
        logger.info(f"Processing course: {course_config['name']}")
        logger.info(f"Course key: {course_key}")
        logger.info(f"Source folders: {course_config['folders']}")
        logger.info(f"{'='*50}")

        # Check which folders exist for this course
        course_folders = []
        for folder_name in course_config['folders']:
            folder_path = data_path / folder_name
            if folder_path.exists() and folder_path.is_dir():
                course_folders.append(folder_path)
            else:
                logger.warning(f"Folder not found: {folder_name}")

        if not course_folders:
            logger.warning(f"No folders found for course {course_config['name']}")
            continue

        # Find all supported documents across all folders for this course
        supported_extensions = {'.pdf', '.docx', '.doc'}
        course_documents = []

        for folder_path in course_folders:
            logger.info(f"Scanning folder: {folder_path.name}")
            for ext in supported_extensions:
                # Look for files in folder and any subdirectories
                docs = list(folder_path.glob(f"**/*{ext}"))
                course_documents.extend(docs)
                logger.info(f"  Found {len(docs)} {ext} files in {folder_path.name}")

        if not course_documents:
            logger.warning(f"No documents found for course {course_config['name']}")
            continue

        logger.info(f"Total documents for {course_config['name']}: {len(course_documents)}")
        for doc in course_documents:
            logger.info(f"  - {doc.name}")

        # Process all documents for this course
        course_chunks = []
        for doc_path in course_documents:
            try:
                logger.info(f"Processing: {doc_path.name}")
                chunks = doc_processor.process_document(
                    file_path=doc_path,
                    course_name=course_config['name'],  # Use the display name
                    professor_name=professor_name
                )
                course_chunks.extend(chunks)
                logger.info(f"  Created {len(chunks)} chunks from {doc_path.name}")

            except Exception as e:
                logger.error(f"Error processing {doc_path}: {e}")
                continue

        if not course_chunks:
            logger.warning(f"No chunks created for course {course_config['name']}")
            continue

        # Generate embeddings for all chunks in this course
        logger.info(f"Generating embeddings for {len(course_chunks)} chunks...")
        course_chunks = doc_processor.generate_embeddings(course_chunks)

        # Upload to Pinecone with course-specific namespace
        namespace = course_key  # Use the course key as namespace
        logger.info(f"Uploading to Pinecone namespace: {namespace}")
        uploader.upload_chunks(course_chunks, namespace)

        total_chunks_processed += len(course_chunks)
        logger.info(f"Successfully processed {len(course_chunks)} chunks for {course_config['name']}")

    # Print final stats
    stats = uploader.get_index_stats()
    logger.info(f"\n{'='*50}")
    logger.info(f"PROCESSING COMPLETE")
    logger.info(f"{'='*50}")
    logger.info(f"Total chunks processed: {total_chunks_processed}")
    logger.info(f"Courses processed: {len(course_mapping)}")
    logger.info(f"Pinecone index stats: {stats}")

def scan_data_directory(data_directory: Path) -> Dict:
    """Scan data directory and return information about found classes and documents"""
    data_path = Path(data_directory)

    if not data_path.exists():
        return {"error": f"Data directory not found: {data_path}"}

    scan_results = {
        "data_directory": str(data_path),
        "classes": {},
        "total_classes": 0,
        "total_documents": 0
    }

    # Find all class folders
    class_folders = [folder for folder in data_path.iterdir()
                    if folder.is_dir() and not folder.name.startswith('.')]

    scan_results["total_classes"] = len(class_folders)

    # Scan each class folder
    supported_extensions = {'.pdf', '.docx', '.doc'}

    for class_folder in class_folders:
        class_info = {
            "folder_name": class_folder.name,
            "folder_path": str(class_folder),
            "documents": [],
            "document_count": 0
        }

        # Find all documents in class folder (including subdirectories)
        for ext in supported_extensions:
            docs = list(class_folder.glob(f"**/*{ext}"))
            for doc in docs:
                rel_path = doc.relative_to(class_folder)
                class_info["documents"].append({
                    "name": doc.name,
                    "relative_path": str(rel_path),
                    "full_path": str(doc),
                    "size_mb": round(doc.stat().st_size / (1024 * 1024), 2),
                    "extension": doc.suffix
                })

        class_info["document_count"] = len(class_info["documents"])
        scan_results["total_documents"] += class_info["document_count"]
        scan_results["classes"][class_folder.name] = class_info

    return scan_results

# Updated main function to work with local file structure
def main():
    """Process all courses from local data directory"""

    # Configuration - UPDATE THESE PATHS AND KEYS
    DATA_DIRECTORY = "documents"  # Path to your documents folder
    PROFESSOR_NAME = "Professor Robert Ceresa"
    PINECONE_API_KEY = "pcsk_UySHG_ErRr5FNDgTKZeC1ZwJSFnjBm8Ggt5aTNZEcJtpuVyYL5ST4No7J9xbWqjVo4UfN"
    PINECONE_INDEX_NAME = "ai-professor-platform"

    print("AI Professor Platform - Document Processing")
    print("=" * 60)

    # First, scan the data directory
    print("Scanning data directory...")
    scan_results = scan_data_directory(Path(DATA_DIRECTORY))

    if "error" in scan_results:
        print(f"Error: {scan_results['error']}")
        print("\nPlease ensure your documents directory structure looks like this:")
        print("documents/")
        print("├── american/")
        print("│   ├── lecture.pdf")
        print("│   └── readings.docx")
        print("├── foundational/")
        print("│   ├── chapter1.docx")
        print("│   └── lecture.pdf")
        print("├── international/")
        print("│   ├── readings.pdf")
        print("│   └── lecture.pdf")
        print("└── theory/")
        print("    └── materials.docx")
        return

    # Display scan results
    print(f"\nFound {scan_results['total_classes']} classes with {scan_results['total_documents']} total documents:")
    print("-" * 60)

    for class_name, class_info in scan_results["classes"].items():
        print(f"📁 {class_name} ({class_info['document_count']} documents)")
        for doc in class_info["documents"][:3]:  # Show first 3 docs
            print(f"   📄 {doc['name']} ({doc['size_mb']} MB)")
        if class_info['document_count'] > 3:
            print(f"   ... and {class_info['document_count'] - 3} more documents")
        print()

    # Confirm before processing
    print("This will:")
    print(f"1. Process {scan_results['total_documents']} documents")
    print(f"2. Create embeddings for all content")
    print(f"3. Upload to Pinecone index: {PINECONE_INDEX_NAME}")
    print(f"4. Create separate namespaces for each class")

    print("\nProceeding with processing...")
    # Auto-proceed without user confirmation

    # Process all courses
    try:
        process_course_content(
            data_directory=Path(DATA_DIRECTORY),
            pinecone_api_key=PINECONE_API_KEY,
            pinecone_index_name=PINECONE_INDEX_NAME,
            professor_name=PROFESSOR_NAME
        )
        print("\n🎉 Processing completed successfully!")
        print("\nNext steps:")
        print("1. Test the Streamlit app locally: streamlit run streamlit_app.py")
        print("2. Deploy to Streamlit Cloud")
        print("3. Share with students")

    except Exception as e:
        logger.error(f"Processing failed: {e}")
        print(f"\n❌ Processing failed: {e}")
        print("Check the logs above for details.")

if __name__ == "__main__":
    main()
import os
import logging
from pathlib import Path
from typing import List, Dict, Optional
import hashlib
from dataclasses import dataclass
from datetime import datetime

# Document processing
from docx import Document
import PyPDF2
from sentence_transformers import SentenceTransformer

# Vector database
from pinecone import Pinecone, ServerlessSpec
import numpy as np

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

@dataclass
class DocumentChunk:
    """Represents a chunk of processed document content"""
    text: str
    metadata: Dict
    chunk_id: str
    embedding: Optional[List[float]] = None

class DocumentProcessor:
    """Handles processing of course documents into embeddings"""

    def __init__(self, embedding_model_name: str = "sentence-transformers/all-mpnet-base-v2"):
        self.embedding_model = SentenceTransformer(embedding_model_name)
        logger.info(f"Initialized embedding model: {embedding_model_name}")

    def extract_text_from_docx(self, file_path: Path) -> str:
        """Extract text from Word document"""
        try:
            doc = Document(file_path)
            full_text = []

            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    full_text.append(paragraph.text.strip())

            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        full_text.append(" | ".join(row_text))

            return "\n\n".join(full_text)

        except Exception as e:
            logger.error(f"Error extracting text from {file_path}: {e}")
            return ""

    def extract_text_from_pdf(self, file_path: Path) -> str:
        """Extract text from PDF document preserving hierarchical structure"""
        try:
            import fitz  # PyMuPDF for better formatting preservation
            text_blocks = []

            with fitz.open(file_path) as pdf_doc:
                for page_num in range(len(pdf_doc)):
                    page = pdf_doc[page_num]

                    # Get text blocks with formatting information
                    blocks = page.get_text("dict")

                    for block in blocks["blocks"]:
                        if "lines" in block:
                            for line in block["lines"]:
                                line_text = ""
                                for span in line["spans"]:
                                    # Preserve indentation by checking x-coordinate
                                    x_coord = span["bbox"][0]
                                    indent = " " * max(0, int((x_coord - 70) / 10))  # Approximate indentation
                                    line_text += indent + span["text"]

                                if line_text.strip():
                                    text_blocks.append(line_text.rstrip())

            return "\n".join(text_blocks)

        except ImportError:
            # Fallback to PyPDF2 if PyMuPDF not available
            try:
                import PyPDF2
                text = []
                with open(file_path, 'rb') as file:
                    pdf_reader = PyPDF2.PdfReader(file)
                    for page in pdf_reader.pages:
                        page_text = page.extract_text()
                        if page_text.strip():
                            text.append(page_text.strip())
                return "\n\n".join(text)
            except Exception as e:
                logger.error(f"Error extracting text from {file_path} with PyPDF2: {e}")
                return ""

        except Exception as e:
            logger.error(f"Error extracting text from {file_path}: {e}")
            return ""

    def extract_text_from_txt(self, file_path: Path) -> str:
        """Extract text from plain text file"""
        try:
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as file:
                return file.read()
        except Exception as e:
            logger.error(f"Error reading text file {file_path}: {e}")
            return ""

    def extract_text_from_file(self, file_path: Path) -> str:
        """Extract text from various file formats"""
        suffix = file_path.suffix.lower()

        if suffix == '.docx':
            return self.extract_text_from_docx(file_path)
        elif suffix == '.pdf':
            return self.extract_text_from_pdf(file_path)
        elif suffix in ['.txt', '.md']:
            return self.extract_text_from_txt(file_path)
        else:
            logger.warning(f"Unsupported file format: {suffix}")
            return ""

    def extract_hierarchical_structure(self, text: str) -> List[Dict]:
        """Extract hierarchical bullet structure from academic content"""
        lines = text.split('\n')
        structured_content = []
        current_lesson = None
        hierarchy_stack = []

        # Bullet symbols in order of hierarchy depth
        bullet_patterns = [
            ('•', 1),  # Main topic bullet
            ('○', 2),  # Subtopic bullet
            ('■', 3),  # Detail bullet
            ('▪', 4),  # Sub-detail bullet
            ('-', 2),  # Alternative subtopic
            ('*', 2),  # Alternative subtopic
        ]

        for line in lines:
            stripped_line = line.strip()
            if not stripped_line:
                continue

            # Detect lesson headers
            if 'lesson' in stripped_line.lower() and any(word in stripped_line.lower() for word in ['lesson', 'chapter', 'module']):
                current_lesson = stripped_line
                hierarchy_stack = [current_lesson]
                continue

            # Detect bullet points and their hierarchy level
            bullet_level = 0
            bullet_symbol = None
            content = stripped_line

            for symbol, level in bullet_patterns:
                if stripped_line.startswith(symbol + ' '):
                    bullet_level = level
                    bullet_symbol = symbol
                    content = stripped_line[len(symbol):].strip()
                    break

            # Calculate indentation level as fallback
            if bullet_level == 0:
                indent_level = len(line) - len(line.lstrip())
                if indent_level > 0:
                    bullet_level = min(4, (indent_level // 2) + 1)

            # Update hierarchy stack
            if bullet_level > 0:
                # Trim stack to current level
                hierarchy_stack = hierarchy_stack[:bullet_level]
                hierarchy_stack.append(content)

                # Create hierarchy path
                hierarchy_path = " > ".join(hierarchy_stack)

                structured_content.append({
                    'content': content,
                    'level': bullet_level,
                    'symbol': bullet_symbol,
                    'hierarchy_path': hierarchy_path,
                    'lesson': current_lesson,
                    'full_line': line
                })
            else:
                # Regular paragraph content
                if hierarchy_stack:
                    hierarchy_path = " > ".join(hierarchy_stack)
                else:
                    hierarchy_path = current_lesson or "General Content"

                structured_content.append({
                    'content': content,
                    'level': 0,
                    'symbol': None,
                    'hierarchy_path': hierarchy_path,
                    'lesson': current_lesson,
                    'full_line': line
                })

        return structured_content

    def hierarchical_chunking(self, structured_content: List[Dict], target_size: int = 900, overlap_size: int = 150) -> List[str]:
        """Create chunks that preserve hierarchical structure and semantic meaning"""
        if not structured_content:
            return []

        chunks = []
        current_chunk_items = []
        current_chunk_size = 0

        for i, item in enumerate(structured_content):
            item_text = item['full_line']
            item_size = len(item_text)

            # Check if adding this item would exceed target size
            if current_chunk_size + item_size > target_size and current_chunk_items:
                # Create chunk from current items
                chunk_text = self._create_chunk_with_context(current_chunk_items)
                chunks.append(chunk_text)

                # Start new chunk with overlap
                overlap_items = self._get_overlap_items(current_chunk_items, overlap_size)
                current_chunk_items = overlap_items + [item]
                current_chunk_size = sum(len(it['full_line']) for it in current_chunk_items)
            else:
                current_chunk_items.append(item)
                current_chunk_size += item_size

            # Force chunk break at lesson boundaries
            if (i < len(structured_content) - 1 and
                item.get('lesson') != structured_content[i + 1].get('lesson') and
                current_chunk_items):
                chunk_text = self._create_chunk_with_context(current_chunk_items)
                chunks.append(chunk_text)
                current_chunk_items = []
                current_chunk_size = 0

        # Add final chunk
        if current_chunk_items:
            chunk_text = self._create_chunk_with_context(current_chunk_items)
            chunks.append(chunk_text)

        return chunks

    def _create_chunk_with_context(self, items: List[Dict]) -> str:
        """Create a chunk with proper hierarchical context"""
        if not items:
            return ""

        # Add lesson context at the top
        lesson = items[0].get('lesson')
        chunk_lines = []

        if lesson:
            chunk_lines.append(f"=== {lesson} ===\n")

        # Group items by hierarchy level for better formatting
        for item in items:
            chunk_lines.append(item['full_line'])

        return '\n'.join(chunk_lines)

    def _get_overlap_items(self, items: List[Dict], overlap_size: int) -> List[Dict]:
        """Get items for overlap, prioritizing higher-level hierarchy items"""
        if not items:
            return []

        # Start from the end and work backwards, prioritizing higher hierarchy levels
        overlap_items = []
        current_size = 0

        for item in reversed(items):
            item_size = len(item['full_line'])
            if current_size + item_size <= overlap_size:
                overlap_items.insert(0, item)
                current_size += item_size
            elif item['level'] <= 2:  # Always include high-level context
                overlap_items.insert(0, item)
                break

        return overlap_items

    def intelligent_chunking(self, text: str, max_chunk_size: int = 900, overlap: int = 150) -> List[str]:
        """Enhanced chunking that preserves hierarchical academic structure"""
        if not text.strip():
            return []

        # First, try hierarchical chunking for structured content
        structured_content = self.extract_hierarchical_structure(text)

        if len(structured_content) > 3:  # If we found good hierarchical structure
            return self.hierarchical_chunking(structured_content, max_chunk_size, overlap)

        # Fallback to paragraph-based chunking for unstructured content
        paragraphs = [p.strip() for p in text.split('\n\n') if p.strip()]

        chunks = []
        current_chunk = ""

        for paragraph in paragraphs:
            # If adding this paragraph would exceed max size, save current chunk
            if len(current_chunk) + len(paragraph) > max_chunk_size and current_chunk:
                chunks.append(current_chunk.strip())

                # Start new chunk with overlap from previous chunk
                if overlap > 0:
                    words = current_chunk.split()
                    overlap_text = " ".join(words[-overlap//4:])  # Last ~25% of words for overlap
                    current_chunk = overlap_text + "\n\n" + paragraph
                else:
                    current_chunk = paragraph
            else:
                if current_chunk:
                    current_chunk += "\n\n" + paragraph
                else:
                    current_chunk = paragraph

        # Add the last chunk
        if current_chunk.strip():
            chunks.append(current_chunk.strip())

        return chunks

    def infer_document_type(self, file_path: Path) -> str:
        """Infer document type from filename and content"""
        filename = file_path.name.lower()

        if 'syllabus' in filename:
            return 'syllabus'
        elif 'lecture' in filename or 'slides' in filename:
            return 'lecture'
        elif 'reading' in filename or 'assignment' in filename:
            return 'reading'
        elif 'exam' in filename or 'quiz' in filename:
            return 'assessment'
        else:
            return 'general'

    def process_document(self, file_path: Path, course_name: str, professor_name: str) -> List[DocumentChunk]:
        """Process a single document into chunks with metadata"""
        logger.info(f"Processing document: {file_path}")

        # Extract text
        text = self.extract_text_from_file(file_path)
        if not text.strip():
            logger.warning(f"No text extracted from {file_path}")
            return []

        # Create chunks using hierarchical chunking
        chunks = self.intelligent_chunking(text)
        logger.info(f"Created {len(chunks)} chunks from {file_path}")

        # Create DocumentChunk objects with enhanced metadata
        document_chunks = []
        for i, chunk_text in enumerate(chunks):
            # Create unique chunk ID
            chunk_content = f"{file_path}_{i}_{chunk_text[:50]}"
            chunk_id = hashlib.md5(chunk_content.encode()).hexdigest()

            # Extract hierarchy information from chunk
            hierarchy_info = self._extract_chunk_hierarchy(chunk_text)

            # Create namespace for consistency with Pinecone storage
            namespace = course_name.lower().replace(" ", "_").replace("-", "_")
            
            metadata = {
                'course': namespace,  # Use namespace for consistency with search filter
                'course_name': course_name,  # Keep original course name for display
                'professor': professor_name,
                'document_name': file_path.name,
                'document_type': self.infer_document_type(file_path),
                'chunk_index': i,
                'total_chunks': len(chunks),
                'file_path': str(file_path),
                'processed_date': datetime.now().isoformat(),
                'text': chunk_text,  # Store text in metadata for Pinecone

                # Hierarchical metadata
                'lesson_number': hierarchy_info.get('lesson_number'),
                'main_topic': hierarchy_info.get('main_topic'),
                'hierarchy_path': hierarchy_info.get('hierarchy_path'),
                'depth_level': hierarchy_info.get('depth_level', 0),
                'chunk_type': hierarchy_info.get('chunk_type', 'general')
            }

            document_chunks.append(DocumentChunk(
                text=chunk_text,
                metadata=metadata,
                chunk_id=chunk_id
            ))

        return document_chunks

    def _extract_chunk_hierarchy(self, chunk_text: str) -> Dict:
        """Extract hierarchical information from a chunk"""
        lines = chunk_text.split('\n')
        hierarchy_info = {
            'lesson_number': None,
            'main_topic': None,
            'hierarchy_path': None,
            'depth_level': 0,
            'chunk_type': 'general'
        }

        for line in lines:
            stripped_line = line.strip()

            # Extract lesson number
            if 'lesson' in stripped_line.lower():
                import re
                lesson_match = re.search(r'lesson\s+(\d+)', stripped_line.lower())
                if lesson_match:
                    hierarchy_info['lesson_number'] = lesson_match.group(1)

            # Detect main topic (usually the first major bullet point)
            if stripped_line.startswith('•') and not hierarchy_info['main_topic']:
                hierarchy_info['main_topic'] = stripped_line[1:].strip()

            # Calculate maximum depth level in chunk
            if any(stripped_line.startswith(symbol) for symbol in ['•', '○', '■', '▪']):
                if stripped_line.startswith('•'):
                    hierarchy_info['depth_level'] = max(hierarchy_info['depth_level'], 1)
                elif stripped_line.startswith('○'):
                    hierarchy_info['depth_level'] = max(hierarchy_info['depth_level'], 2)
                elif stripped_line.startswith('■'):
                    hierarchy_info['depth_level'] = max(hierarchy_info['depth_level'], 3)
                elif stripped_line.startswith('▪'):
                    hierarchy_info['depth_level'] = max(hierarchy_info['depth_level'], 4)

        # Determine chunk type
        if hierarchy_info['lesson_number']:
            hierarchy_info['chunk_type'] = 'lesson_content'
        elif hierarchy_info['depth_level'] > 0:
            hierarchy_info['chunk_type'] = 'structured_content'

        # Create hierarchy path
        path_components = []
        if hierarchy_info['lesson_number']:
            path_components.append(f"Lesson {hierarchy_info['lesson_number']}")
        if hierarchy_info['main_topic']:
            path_components.append(hierarchy_info['main_topic'][:50])  # Truncate if too long

        hierarchy_info['hierarchy_path'] = " > ".join(path_components) if path_components else "General Content"

        return hierarchy_info

    def generate_embeddings(self, chunks: List[DocumentChunk]) -> List[DocumentChunk]:
        """Generate embeddings for document chunks"""
        if not chunks:
            return chunks

        logger.info(f"Generating embeddings for {len(chunks)} chunks")

        # Extract text for embedding
        texts = [chunk.text for chunk in chunks]

        # Generate embeddings
        embeddings = self.embedding_model.encode(texts, show_progress_bar=True)

        # Add embeddings to chunks
        for chunk, embedding in zip(chunks, embeddings):
            chunk.embedding = embedding.tolist()

        return chunks

class PineconeUploader:
    """Handles uploading processed documents to Pinecone"""

    def __init__(self, api_key: str, index_name: str, dimension: int = 768):
        self.pc = Pinecone(api_key=api_key)
        self.index_name = index_name
        self.dimension = dimension

        # Create index if it doesn't exist
        self._ensure_index_exists()
        self.index = self.pc.Index(index_name)
        logger.info(f"Connected to Pinecone index: {index_name}")

    def _ensure_index_exists(self):
        """Create Pinecone index if it doesn't exist"""
        existing_indexes = [index.name for index in self.pc.list_indexes()]

        if self.index_name not in existing_indexes:
            logger.info(f"Creating new Pinecone index: {self.index_name}")
            self.pc.create_index(
                name=self.index_name,
                dimension=self.dimension,
                metric="cosine",
                spec=ServerlessSpec(cloud="aws", region="us-east-1")
            )
        else:
            logger.info(f"Using existing Pinecone index: {self.index_name}")

    def upload_chunks(self, chunks: List[DocumentChunk], namespace: str, batch_size: int = 100):
        """Upload chunks to Pinecone with specified namespace"""
        if not chunks:
            logger.warning("No chunks to upload")
            return

        logger.info(f"Uploading {len(chunks)} chunks to namespace: {namespace}")

        # Prepare vectors for upsert
        vectors = []
        for chunk in chunks:
            if chunk.embedding is None:
                logger.warning(f"Skipping chunk without embedding: {chunk.chunk_id}")
                continue

            vectors.append({
                "id": chunk.chunk_id,
                "values": chunk.embedding,
                "metadata": chunk.metadata
            })

        # Upload in batches
        for i in range(0, len(vectors), batch_size):
            batch = vectors[i:i + batch_size]
            try:
                self.index.upsert(vectors=batch, namespace=namespace)
                logger.info(f"Uploaded batch {i//batch_size + 1}/{(len(vectors)-1)//batch_size + 1}")
            except Exception as e:
                logger.error(f"Error uploading batch {i//batch_size + 1}: {e}")

        logger.info(f"Successfully uploaded {len(vectors)} chunks to namespace: {namespace}")

    def get_index_stats(self) -> Dict:
        """Get statistics about the Pinecone index"""
        return self.index.describe_index_stats()

def process_course_content(
    data_directory: Path,
    pinecone_api_key: str,
    pinecone_index_name: str,
    professor_name: str = "Professor"
):
    """Main function to process all courses from local data directory structure"""
    data_path = Path(data_directory)

    if not data_path.exists():
        logger.error(f"Data directory not found: {data_path}")
        return

    logger.info(f"Starting processing from data directory: {data_path}")

    # Initialize processors
    doc_processor = DocumentProcessor()
    uploader = PineconeUploader(pinecone_api_key, pinecone_index_name)

    # Find all class folders in data directory
    class_folders = [folder for folder in data_path.iterdir()
                    if folder.is_dir() and not folder.name.startswith('.')]

    if not class_folders:
        logger.warning(f"No class folders found in {data_path}")
        return

    logger.info(f"Found {len(class_folders)} class folders: {[f.name for f in class_folders]}")

    # Process each class folder
    total_chunks_processed = 0

    for class_folder in class_folders:
        logger.info(f"\n{'='*50}")
        logger.info(f"Processing class: {class_folder.name}")
        logger.info(f"Directory: {class_folder}")
        logger.info(f"{'='*50}")

        # Find all supported documents in this class folder
        supported_extensions = {'.pdf', '.docx', '.doc'}
        class_documents = []

        for ext in supported_extensions:
            # Look for files in class folder and any subdirectories
            class_documents.extend(class_folder.glob(f"**/*{ext}"))

        if not class_documents:
            logger.warning(f"No documents found in {class_folder}")
            continue

        logger.info(f"Found {len(class_documents)} documents in {class_folder.name}")
        for doc in class_documents:
            logger.info(f"  - {doc.name}")

        # Process all documents for this class
        class_chunks = []
        for doc_path in class_documents:
            try:
                logger.info(f"Processing: {doc_path.name}")
                chunks = doc_processor.process_document(
                    file_path=doc_path,
                    course_name=class_folder.name,
                    professor_name=professor_name
                )
                class_chunks.extend(chunks)
                logger.info(f"  Created {len(chunks)} chunks from {doc_path.name}")

            except Exception as e:
                logger.error(f"Error processing {doc_path}: {e}")
                continue

        if not class_chunks:
            logger.warning(f"No chunks created for class {class_folder.name}")
            continue

        # Generate embeddings for all chunks in this class
        logger.info(f"Generating embeddings for {len(class_chunks)} chunks...")
        class_chunks = doc_processor.generate_embeddings(class_chunks)

        # Upload to Pinecone with class-specific namespace
        namespace = class_folder.name.lower().replace(" ", "_").replace("-", "_")
        logger.info(f"Uploading to Pinecone namespace: {namespace}")
        uploader.upload_chunks(class_chunks, namespace)

        total_chunks_processed += len(class_chunks)
        logger.info(f"Successfully processed {len(class_chunks)} chunks for {class_folder.name}")

    # Print final stats
    stats = uploader.get_index_stats()
    logger.info(f"\n{'='*50}")
    logger.info(f"PROCESSING COMPLETE")
    logger.info(f"{'='*50}")
    logger.info(f"Total chunks processed: {total_chunks_processed}")
    logger.info(f"Classes processed: {len(class_folders)}")
    logger.info(f"Pinecone index stats: {stats}")

def scan_data_directory(data_directory: Path) -> Dict:
    """Scan data directory and return information about found classes and documents"""
    data_path = Path(data_directory)

    if not data_path.exists():
        return {"error": f"Data directory not found: {data_path}"}

    scan_results = {
        "data_directory": str(data_path),
        "classes": {},
        "total_classes": 0,
        "total_documents": 0
    }

    # Find all class folders
    class_folders = [folder for folder in data_path.iterdir()
                    if folder.is_dir() and not folder.name.startswith('.')]

    scan_results["total_classes"] = len(class_folders)

    # Scan each class folder
    supported_extensions = {'.pdf', '.docx', '.doc'}

    for class_folder in class_folders:
        class_info = {
            "folder_name": class_folder.name,
            "folder_path": str(class_folder),
            "documents": [],
            "document_count": 0
        }

        # Find all documents in class folder (including subdirectories)
        for ext in supported_extensions:
            docs = list(class_folder.glob(f"**/*{ext}"))
            for doc in docs:
                rel_path = doc.relative_to(class_folder)
                class_info["documents"].append({
                    "name": doc.name,
                    "relative_path": str(rel_path),
                    "full_path": str(doc),
                    "size_mb": round(doc.stat().st_size / (1024 * 1024), 2),
                    "extension": doc.suffix
                })

        class_info["document_count"] = len(class_info["documents"])
        scan_results["total_documents"] += class_info["document_count"]
        scan_results["classes"][class_folder.name] = class_info

    return scan_results

# Updated main function to work with local file structure
def main():
    """Process all courses from local data directory"""

    # Configuration - UPDATE THESE PATHS AND KEYS
    DATA_DIRECTORY = "documents"  # Path to your documents folder
    PROFESSOR_NAME = "Professor [Father's Name]"
    PINECONE_API_KEY = "your-pinecone-api-key"
    PINECONE_INDEX_NAME = "ai-professor-platform"

    print("AI Professor Platform - Document Processing")
    print("=" * 60)

    # First, scan the data directory
    print("Scanning data directory...")
    scan_results = scan_data_directory(Path(DATA_DIRECTORY))

    if "error" in scan_results:
        print(f"Error: {scan_results['error']}")
        print("\nPlease ensure your documents directory structure looks like this:")
        print("documents/")
        print("├── american/")
        print("│   ├── lecture.pdf")
        print("│   └── readings.docx")
        print("├── foundational/")
        print("│   ├── chapter1.docx")
        print("│   └── lecture.pdf")
        print("├── international/")
        print("│   ├── readings.pdf")
        print("│   └── lecture.pdf")
        print("└── theory/")
        print("    └── materials.docx")
        return

    # Display scan results
    print(f"\nFound {scan_results['total_classes']} classes with {scan_results['total_documents']} total documents:")
    print("-" * 60)

    for class_name, class_info in scan_results["classes"].items():
        print(f"📁 {class_name} ({class_info['document_count']} documents)")
        for doc in class_info["documents"][:3]:  # Show first 3 docs
            print(f"   📄 {doc['name']} ({doc['size_mb']} MB)")
        if class_info['document_count'] > 3:
            print(f"   ... and {class_info['document_count'] - 3} more documents")
        print()

    # Confirm before processing
    print("This will:")
    print(f"1. Process {scan_results['total_documents']} documents")
    print(f"2. Create embeddings for all content")
    print(f"3. Upload to Pinecone index: {PINECONE_INDEX_NAME}")
    print(f"4. Create separate namespaces for each class")

    response = input("\nProceed with processing? (y/N): ")
    if response.lower() != 'y':
        print("Processing cancelled.")
        return

    # Process all courses
    try:
        process_course_content(
            data_directory=Path(DATA_DIRECTORY),
            pinecone_api_key=PINECONE_API_KEY,
            pinecone_index_name=PINECONE_INDEX_NAME,
            professor_name=PROFESSOR_NAME
        )
        print("\n🎉 Processing completed successfully!")
        print("\nNext steps:")
        print("1. Test the Streamlit app locally: streamlit run streamlit_app.py")
        print("2. Deploy to Streamlit Cloud")
        print("3. Share with students")

    except Exception as e:
        logger.error(f"Processing failed: {e}")
        print(f"\n❌ Processing failed: {e}")
        print("Check the logs above for details.")